from docx import Document
from openpyxl import load_workbook

doc_path = r"D:\FULearning\LAB211\J1.S.P0009.docx"
xlsx_path = r"D:\FULearning\LAB211\checklist_review_code_java_vi.xlsx"

doc = Document(doc_path)
print("=== ASSIGNMENT ===")
for i, paragraph in enumerate(doc.paragraphs, 1):
    value = paragraph.text.strip()
    if value:
        print(f"P{i}: {value}")
for table_index, table in enumerate(doc.tables, 1):
    print(f"TABLE {table_index}")
    for row in table.rows:
        print(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells))

wb = load_workbook(xlsx_path, data_only=False)
print("=== CHECKLIST ===")
for ws in wb.worksheets:
    print(f"SHEET: {ws.title}; used={ws.max_row}x{ws.max_column}")
    for row in ws.iter_rows():
        values = [str(cell.value).strip() if cell.value is not None else "" for cell in row]
        if any(values):
            print(" | ".join(values))
